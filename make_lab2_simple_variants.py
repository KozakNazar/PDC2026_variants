##########################################
# N.Kozak // Lviv'2026                   #
#    file: make_lab2_simple_variuants.py #
##########################################
from docx import Document
from docx.shared import Pt, Cm #, Inches, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

from make_labs_common_functions import replace_skip, convert_to_math_unicode_formula, parse_variants_file, lab2_variants_redef

from docx.shared import Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_column_width(column, width):
    """Встановити ширину колонки через XML"""
    # Встановлюємо ширину для кожної комірки в колонці
    for cell in column.cells:
        cell.width = width
        
        # Додатково встановлюємо через XML
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(int(width.twips)))  # twips - 1/20 пункту
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)

def make_lab2_simple_variants_and_save_in_docx(variants, year, group, output_file):
    doc = Document()

    # Заголовок
    title = doc.add_paragraph("ВАРІАНТИ ЗАВДАНЬ (КІ-" + str(group) + ", " + str(year - 1) + "/" + str(year) + " н.р.)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    doc.add_paragraph(
        "Задано (вхідні дані):\n"
        "y₁ і y₂ - вектори-стовпці з n елементів,\n"
        "Y₃ - квадратна матриця порядку n,\n"
        "yᵀ означає операцію транспонування для y."
    )
    
    table = doc.add_table(rows=len(variants), cols=2)
    table.style = 'Table Grid'
    
    # Заповнюємо дані
    for i, variant in enumerate(variants):
        row = table.rows[i]
   
        cell_num = row.cells[0]
        cell_num.text = str(variant['number'])
        cell_num.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        cell_formula = row.cells[1]
        cell_formula.text = convert_to_math_unicode_formula(variant['formula'], True) + ", " + variant['type']
        cell_formula.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Встановлюємо ширину колонок
    set_column_width(table.columns[0], Cm(1.2))  # Мінімальна для номера
    set_column_width(table.columns[1], Cm(14.8)) # Решта місця (для A4)
    
    # АБО можна зробити пропорційно:
    # page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
    # table.columns[1].width = page_width - Cm(1.5)
    
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
                    run.font.name = 'Times New Roman'
    
    doc.save(output_file)
    print("Created:", output_file)

if __name__ == "__main__":
    YEAR_FIRST = 2026
    YEAR_LAST = 2027
    GROUP_FIRST = 301
    GROUP_LAST = 309
    LAB2_VARIANTS_DATA_FILE_NAME = "lab2_simple_variants_data.txt"

    # Читаємо базові варіанти
    variants_ = parse_variants_file(LAB2_VARIANTS_DATA_FILE_NAME, True)

    for year in range(YEAR_FIRST, YEAR_LAST + 1):
        for group in range(GROUP_FIRST, GROUP_LAST + 1):
            # Створюємо нові варіанти
            variants = variants_#lab2_variants_redef(variants_, year, group, True)
            
            # Створюємо шлях до файлу
            filename = "PRO_LAB2_VARIANTSANDMATLAB/" + str(year - 1) + "_" + str(year) + "/KI" + str(group) + "/l2_variants_ki" + str(group) + "_" + str(year - 1) + str(year) + ".docx"
            path = Path(filename)            
            path.parent.mkdir(parents=True, exist_ok=True) # if path.parent != Path("."):

            make_lab2_simple_variants_and_save_in_docx(
                variants,
                year,
                group, 
                filename                         
            )